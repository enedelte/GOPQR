from docx import Document
document = Document()
paragraph1 = document.add_paragraph('Bogotá,(insetar fecha hoy)')
paragraph2 = document.add_paragraph('VO-GA_DGO-(insertar radicado)-23')
paragraph3 = document.add_paragraph('Señor(es)')
paragraph4 = document.add_paragraph('(insertar nombre reclamante)')
paragraph5 = document.add_paragraph('(insertar tipo documento reclamante): (insertar documento reclamante)')
paragraph6 = document.add_paragraph('Dirección: (insertar dirección)')
paragraph7 = document.add_paragraph('Teléfono: (insertar teléfono)')
paragraph8 = document.add_paragraph('Correo: (insertar correo)')
paragraph9 = document.add_paragraph('(insertar municipio/departamento)')
paragraph10 = document.add_paragraph('Respetado(a) señor(a): Reciba un cordial saludo en nombre de NUEVA EPS S.A. Agradecemos su confianza al exponernos sus inquietudes, lo que nos permite trabajar permanentemente y así identificar acciones de mejora que conlleven a fortalecer nuestro servicio.')
paragraph11 = document.add_paragraph('Porque nos interesa ofrecerle un mejor servicio, queremos que nos cuente cómo fue su experiencia con la solución de su solicitud PQRS para lo cual lo invitamos a contestar dos preguntas en el siguiente enlace https://forms.office.com/r/pJaFmjkLW1')
Paragraph12 = document.add_paragraph('Esperamos haber aclarado su inquietud y le expresamos nuestra permanente disposición para atenderlo (a). Recuerde que NUEVA EPS S.A., tiene a su alcance varios canales de atención al usuario, para aclararle cualquier inquietud o suministrarle la información que usted requiera a través de la línea en Bogotá 601 3077022 y la línea gratuita 018000954400 para el resto del país, nuestra APP Nueva EPS o a través de nuestro portal web www.nuevaeps.com.co.')
paragraph13 = document.add_paragraph('Cordialmente,')
## resolver document.add_picture('C:\Users\daale\OneDrive - NUEVA EPS\GO_PQR\querys\firmajc.png')
paragraph14 = document.add_paragraph('DIRECCION DE GESTION OPERATIVA')
paragraph15 = document.add_paragraph('Gerencia de Afiliaciones')
paragraph16 = document.add_paragraph('Vicepresidencia de Operaciones')
paragraph17 = document.add_paragraph('Nueva EPS')
paragraph18 = document.add_paragraph('Elaboró: (insertad usuario)')
paragraph19 = document.add_paragraph("“Frente a cualquier desacuerdo en la decisión adoptada por la EPS a la cual se elevó la respectiva queja o petición, se puede elevar consulta ante la correspondiente Dirección de Salud, sea esta la Departamental, Distrital o Local, sin perjuicio de la competencia prevalente y excluyente que le corresponde a la Superintendencia Nacional de Salud, como autoridad máxima en materia de inspección, vigilancia y control de este sector”.")

document.save('micartaauto.docx')

